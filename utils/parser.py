import os
import re
import fitz
import docx
import pytesseract
from PIL import Image

ALLOWED_EXTENSIONS = {'pdf', 'docx', 'txt', 'png', 'jpg', 'jpeg'}
ALLOWED_MIME_TYPES = {
    'application/pdf',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain',
    'image/png',
    'image/jpeg'
}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(file_path):
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
    except Exception as e:
        raise Exception(f"Error extracting PDF: {str(e)}")
    return text.strip()

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {str(e)}")
    return text.strip()

def extract_text_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read().strip()
        except Exception as e2:
            raise Exception(f"Error extracting TXT: {str(e2)}")

def extract_text_from_image(file_path):
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from image: {str(e)}")

def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower().replace('.', '')
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif ext == 'docx':
        return extract_text_from_docx(file_path)
    elif ext == 'txt':
        return extract_text_from_txt(file_path)
    elif ext in ['png', 'jpg', 'jpeg']:
        return extract_text_from_image(file_path)
    else:
        raise Exception(f"Unsupported file format: {ext}")

def extract_contact_info(text):
    info = {}
    
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    info['email'] = emails[0] if emails else None
    
    phone_pattern = r'(?:\+?1[-.\s]?)?\(?[0-9]{3}\)?[-.\s]?[0-9]{3}[-.\s]?[0-9]{4}'
    phones = re.findall(phone_pattern, text)
    info['phone'] = phones[0] if phones else None
    
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
    info['linkedin'] = linkedin[0] if linkedin else None
    
    github_pattern = r'github\.com/[\w-]+'
    github = re.findall(github_pattern, text, re.IGNORECASE)
    info['github'] = github[0] if github else None
    
    return info

def extract_name(text):
    lines = text.strip().split('\n')
    if lines:
        first_line = lines[0].strip()
        if len(first_line.split()) <= 4 and len(first_line) < 50:
            if not '@' in first_line and not re.match(r'^\d', first_line):
                return first_line
    return None

def parse_resume(file_path):
    text = extract_text(file_path)
    
    if not text:
        raise Exception("No text could be extracted from the file")
    
    result = {
        'text': text,
        'name': extract_name(text),
        'contact': extract_contact_info(text),
        'word_count': len(text.split()),
        'char_count': len(text)
    }
    
    return result